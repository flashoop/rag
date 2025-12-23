#!/usr/bin/env python3
"""
Word (DOCX) to Markdown Batch Converter

Converts Microsoft Word (.docx) files to Markdown format using multiple methods:
1. pandoc - Best quality, requires system installation (recommended)
2. python-docx + markdownify - Pure Python solution
3. mammoth - Good for complex formatting

Usage:
    python word_to_markdown.py input_folder output_folder
    python word_to_markdown.py input_folder output_folder --method pandoc
    python word_to_markdown.py single_file.docx output.md
"""

import os
import sys
import argparse
import subprocess
from pathlib import Path
from typing import Optional, List
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class WordToMarkdownConverter:
    """Converts Word documents to Markdown using multiple methods"""

    def __init__(self, method: str = 'auto'):
        """
        Initialize converter with specified method

        Args:
            method: 'pandoc', 'python-docx', 'mammoth', or 'auto'
        """
        self.method = method
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which conversion methods are available"""
        self.available_methods = []

        # Check for pandoc (system installation)
        try:
            result = subprocess.run(
                ['pandoc', '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode == 0:
                self.available_methods.append('pandoc')
                logger.info("✓ pandoc available")
        except (FileNotFoundError, subprocess.TimeoutExpired):
            logger.debug("✗ pandoc not available")

        # Check for python-docx
        try:
            import docx
            import markdownify
            self.available_methods.append('python-docx')
            logger.info("✓ python-docx + markdownify available")
        except ImportError:
            logger.debug("✗ python-docx + markdownify not available")

        # Check for mammoth
        try:
            import mammoth
            self.available_methods.append('mammoth')
            logger.info("✓ mammoth available")
        except ImportError:
            logger.debug("✗ mammoth not available")

        if not self.available_methods:
            raise ImportError(
                "No Word conversion libraries found. Install at least one:\n"
                "  System: Install pandoc (https://pandoc.org/installing.html)\n"
                "  Python: pip install python-docx markdownify\n"
                "  Python: pip install mammoth"
            )

    def convert_with_pandoc(self, docx_path: str) -> str:
        """
        Convert DOCX to Markdown using pandoc (best quality)

        Args:
            docx_path: Path to DOCX file

        Returns:
            Markdown content as string
        """
        try:
            result = subprocess.run(
                [
                    'pandoc',
                    docx_path,
                    '-f', 'docx',
                    '-t', 'markdown',
                    '--wrap=none',  # Don't wrap lines
                    '--extract-media=.',  # Extract images
                ],
                capture_output=True,
                text=True,
                check=True,
                timeout=60
            )
            return result.stdout

        except subprocess.CalledProcessError as e:
            logger.error(f"pandoc failed: {e.stderr}")
            raise
        except subprocess.TimeoutExpired:
            logger.error("pandoc timed out")
            raise

    def convert_with_python_docx(self, docx_path: str) -> str:
        """
        Convert DOCX to Markdown using python-docx

        Args:
            docx_path: Path to DOCX file

        Returns:
            Markdown content as string
        """
        from docx import Document
        from markdownify import markdownify as md

        try:
            doc = Document(docx_path)
            markdown_parts = []

            # Process paragraphs
            for para in doc.paragraphs:
                # Get paragraph text
                text = para.text.strip()
                if not text:
                    markdown_parts.append('')
                    continue

                # Detect heading style
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.split()[-1])
                    markdown_parts.append(f"{'#' * level} {text}")
                elif para.style.name == 'List Bullet':
                    markdown_parts.append(f"- {text}")
                elif para.style.name.startswith('List Number'):
                    markdown_parts.append(f"1. {text}")
                else:
                    # Regular paragraph with inline formatting
                    formatted_text = self._format_runs(para)
                    markdown_parts.append(formatted_text)

            # Process tables
            for table in doc.tables:
                markdown_table = self._format_table(table)
                markdown_parts.append(markdown_table)

            return '\n\n'.join(markdown_parts)

        except Exception as e:
            logger.error(f"python-docx failed: {e}")
            raise

    def _format_runs(self, paragraph) -> str:
        """Format runs (text segments) with inline markdown"""
        formatted_parts = []

        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            if run.font.strike:
                text = f"~~{text}~~"

            formatted_parts.append(text)

        return ''.join(formatted_parts)

    def _format_table(self, table) -> str:
        """Format table as Markdown table"""
        rows = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        # Format header
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "|" + "|".join([" --- " for _ in rows[0]]) + "|"

        # Format data rows
        data_rows = []
        for row in rows[1:]:
            data_rows.append("| " + " | ".join(row) + " |")

        return "\n".join([header, separator] + data_rows)

    def convert_with_mammoth(self, docx_path: str) -> str:
        """
        Convert DOCX to Markdown using mammoth

        Args:
            docx_path: Path to DOCX file

        Returns:
            Markdown content as string
        """
        import mammoth

        try:
            with open(docx_path, 'rb') as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                markdown = result.value

                # Log warnings if any
                if result.messages:
                    for message in result.messages:
                        logger.warning(f"mammoth: {message}")

                return markdown

        except Exception as e:
            logger.error(f"mammoth failed: {e}")
            raise

    def convert(self, docx_path: str) -> str:
        """
        Convert DOCX to Markdown using best available method

        Args:
            docx_path: Path to DOCX file

        Returns:
            Markdown content as string
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word file not found: {docx_path}")

        logger.info(f"Converting: {docx_path}")

        # Determine method to use
        if self.method == 'auto':
            # Try methods in order of quality
            methods = ['pandoc', 'mammoth', 'python-docx']
            methods = [m for m in methods if m in self.available_methods]
        else:
            if self.method not in self.available_methods:
                raise ValueError(
                    f"Method '{self.method}' not available. "
                    f"Available: {self.available_methods}"
                )
            methods = [self.method]

        # Try each method until one succeeds
        last_error = None
        for method in methods:
            try:
                logger.info(f"Trying method: {method}")

                if method == 'pandoc':
                    return self.convert_with_pandoc(docx_path)
                elif method == 'python-docx':
                    return self.convert_with_python_docx(docx_path)
                elif method == 'mammoth':
                    return self.convert_with_mammoth(docx_path)

            except Exception as e:
                last_error = e
                logger.warning(f"Method {method} failed: {e}")
                continue

        # All methods failed
        raise RuntimeError(
            f"All conversion methods failed. Last error: {last_error}"
        )

    def convert_file(
        self,
        docx_path: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        Convert single Word file to Markdown

        Args:
            docx_path: Path to Word file
            output_path: Path to output Markdown file (optional)

        Returns:
            Path to output file
        """
        # Convert to Markdown
        markdown_content = self.convert(docx_path)

        # Determine output path
        if output_path is None:
            output_path = Path(docx_path).with_suffix('.md')
        else:
            output_path = Path(output_path)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            # Add metadata header
            f.write(f"# {Path(docx_path).stem}\n\n")
            f.write(f"*Converted from: {Path(docx_path).name}*\n\n")
            f.write("---\n\n")
            f.write(markdown_content)

        logger.info(f"✓ Saved: {output_path}")
        return str(output_path)

    def convert_batch(
        self,
        input_dir: str,
        output_dir: Optional[str] = None,
        recursive: bool = False
    ) -> List[str]:
        """
        Batch convert Word files in directory

        Args:
            input_dir: Directory containing Word files
            output_dir: Directory for output Markdown files
            recursive: Whether to process subdirectories

        Returns:
            List of output file paths
        """
        input_path = Path(input_dir)

        if not input_path.exists():
            raise FileNotFoundError(f"Input directory not found: {input_dir}")

        # Find Word files (.docx and .doc)
        if recursive:
            docx_files = list(input_path.rglob("*.docx"))
            docx_files.extend(list(input_path.rglob("*.doc")))
        else:
            docx_files = list(input_path.glob("*.docx"))
            docx_files.extend(list(input_path.glob("*.doc")))

        # Filter out temporary files (starting with ~$)
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]

        if not docx_files:
            logger.warning(f"No Word files found in {input_dir}")
            return []

        logger.info(f"Found {len(docx_files)} Word file(s)")

        # Determine output directory
        if output_dir is None:
            output_dir = input_path
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

        # Convert each file
        output_files = []
        success_count = 0
        fail_count = 0

        for docx_file in docx_files:
            try:
                # Preserve directory structure if recursive
                if recursive:
                    relative_path = docx_file.relative_to(input_path)
                    output_path = output_dir / relative_path.with_suffix('.md')
                else:
                    output_path = output_dir / docx_file.with_suffix('.md').name

                # Convert file
                result = self.convert_file(str(docx_file), str(output_path))
                output_files.append(result)
                success_count += 1

            except Exception as e:
                logger.error(f"✗ Failed to convert {docx_file}: {e}")
                fail_count += 1

        # Print summary
        logger.info("\n" + "=" * 60)
        logger.info(f"Conversion Complete!")
        logger.info(f"  Success: {success_count}/{len(docx_files)}")
        logger.info(f"  Failed:  {fail_count}/{len(docx_files)}")
        logger.info("=" * 60)

        return output_files


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description='Batch convert Word files to Markdown format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert single file
  python word_to_markdown.py input.docx output.md

  # Convert all Word files in a folder
  python word_to_markdown.py input_folder/ output_folder/

  # Convert with specific method
  python word_to_markdown.py input_folder/ output_folder/ --method pandoc

  # Convert recursively
  python word_to_markdown.py input_folder/ output_folder/ --recursive

Available methods:
  - pandoc: Best quality, requires system installation (recommended)
  - mammoth: Good for complex formatting
  - python-docx: Pure Python, handles tables well
  - auto: Try all methods (default)

Installation:
  System: Install pandoc from https://pandoc.org/installing.html
  Python: pip install python-docx markdownify mammoth
        """
    )

    parser.add_argument(
        'input',
        help='Input Word file or directory'
    )
    parser.add_argument(
        'output',
        nargs='?',
        default=None,
        help='Output Markdown file or directory (default: same as input)'
    )
    parser.add_argument(
        '--method',
        choices=['auto', 'pandoc', 'python-docx', 'mammoth'],
        default='auto',
        help='Conversion method to use (default: auto)'
    )
    parser.add_argument(
        '--recursive',
        '-r',
        action='store_true',
        help='Process subdirectories recursively'
    )
    parser.add_argument(
        '--verbose',
        '-v',
        action='store_true',
        help='Enable verbose logging'
    )

    args = parser.parse_args()

    # Set log level
    if args.verbose:
        logger.setLevel(logging.DEBUG)

    # Initialize converter
    try:
        converter = WordToMarkdownConverter(method=args.method)
    except ImportError as e:
        logger.error(str(e))
        sys.exit(1)

    # Convert files
    try:
        input_path = Path(args.input)

        if input_path.is_file():
            # Single file conversion
            converter.convert_file(args.input, args.output)

        elif input_path.is_dir():
            # Batch conversion
            converter.convert_batch(
                args.input,
                args.output,
                recursive=args.recursive
            )

        else:
            logger.error(f"Invalid input path: {args.input}")
            sys.exit(1)

    except Exception as e:
        logger.error(f"Conversion failed: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
